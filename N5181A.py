import pyvisa
import time
import os
from docx import Document
from datetime import datetime

# Kết nối hai máy đo

rm = pyvisa.ResourceManager('@py')

ip_1 = "TCPIP0::169.254.207.61::inst0::INSTR"
ip_2 = "TCPIP0::169.254.207.112::inst0::INSTR"

sig_anl = rm.open_resource(ip_1)
sig_gen = rm.open_resource(ip_2)
try:
	sig_anl = rm.open_resource(ip_1)
	print("✅ Kết nối thành công!")
	print("Thông tin máy đo:", sig_anl.query("*IDN?"))
except pyvisa.VisaIOError as e:
	print("❌ Lỗi kết nối:", e)
try:
	sig_gen = rm.open_resource(ip_2)
	print("✅ Kết nối thành công!")
	print("Thông tin máy đo:", sig_gen.query("*IDN?"))
except pyvisa.VisaIOError as e:
	print("❌ Lỗi kết nối:", e)

# Kiểm tra tần số
print("🔹 Kiểm tra tần số")
sig_anl.write("SYST:PRES:USER:ALL")
time.sleep(1)
sig_gen.write("SYST:PRES")
time.sleep(1)
freq_list = [0.1, 200, 500, 2000, 3000]
results_freq = []

for freq in freq_list:
	if freq < 11 :
		sig_gen.write(f"FREQ {freq} MHz")
		sig_gen.write(f"POW -10 dBm")
		sig_gen.write("OUTP ON")
		time.sleep(1)

		sig_anl.write("FREQ:START 0 Hz")
		sig_anl.write("FREQ:STOP 1 MHz")
		sig_anl.write("BAND 5 kHz")  
		sig_anl.write("BAND:VID 1 kHz")
		time.sleep(1)
		sig_anl.write("CALC:MARK:MAX")
		sig_anl.write(f"SENS:FREQ:CENT {sig_anl.query('CALC:MARK:X?')}")
		time.sleep(1)

		sig_anl.write("FREQ:SPAN  100 kHz")
		sig_anl.write("BAND 100 Hz")  
		sig_anl.write("BAND:VID 100 Hz")
		sig_anl.write("CALC:MARK:MAX")
		sig_anl.write(f"SENS:FREQ:CENT {sig_anl.query('CALC:MARK:X?')}")
		time.sleep(1)
		
		sig_anl.write("FREQ:SPAN  10 kHz")
		sig_anl.write("BAND 10 Hz")  
		sig_anl.write("BAND:VID 10 Hz")
		sig_anl.write("CALC:MARK:MAX")
		sig_anl.write(f"SENS:FREQ:CENT {sig_anl.query('CALC:MARK:X?')}")
		time.sleep(1)

		sig_anl.write("FREQ:SPAN 10 Hz")
		sig_anl.write("BAND 1 Hz")  
		sig_anl.write("BAND:VID 1 Hz")
		time.sleep(2)

		sig_anl.write("CALC:MARK:MAX")
		time.sleep(4)
		sig_anl.write("CALC:MARK:MAX")
		

		peak_freq = round(float(sig_anl.query("CALC:MARK:X?")), 4)
		peak_ampl = float(sig_anl.query("CALC:MARK:Y?"))
		freq_error = format(0.1 - (peak_freq / 1000000), ".12f")
		results_freq.append((freq,peak_freq,freq_error))
	else :
		sig_gen.write(f"FREQ {freq} MHz")
		sig_gen.write(f"POW -10 dBm")
		sig_gen.write("OUTP ON")
		time.sleep(2)
		sig_anl.write("CALC:MARK:MAX") 
		sig_anl.write("SENS:FREQ:SPAN MAX")
		sig_anl.write("SENS:FREQ:TUNE:IMM")
		time.sleep(1)
		peak_freq = float(sig_anl.query("CALC:MARK:X?"))
		freq_error=freq*1000000-peak_freq
		results_freq.append((freq,peak_freq,freq_error))
print("\n📊 Kết quả kiểm tra tần số:")
print("Tần số đặt (MHz) | Tần số đo (KHz) | Sai số (Hz)")
print("-" * 50)
for f_set, f_meas, f_err in results_freq:
    print(f"{f_set:<16} | {f_meas:<16} | {f_err:<16}")

# Kiểm tra biên độ máy phát
print("🔹 Kiểm tra biên độ")
sig_gen.write("SYST:PRES")
time.sleep(1)
sig_anl.write("SYST:PRES:USER:ALL")
time.sleep(1)

amp_list = [125, 275, 425, 850, 1075, 2700, 3000]
results_amp = []
for amp in amp_list:
		sig_gen.write(f"FREQ {freq} MHz")
		sig_gen.write(f"POW 7 dBm")
		sig_gen.write("OUTP ON")
		time.sleep(2)
		sig_anl.write("CALC:MARK:MAX") 
		sig_anl.write("SENS:FREQ:SPAN MAX")
		sig_anl.write("SENS:FREQ:TUNE:IMM")
		time.sleep(1)
		peak_amp = float(sig_anl.query("CALC:MARK:Y?"))
		amp_error = 7 - peak_amp
		results_amp.append((amp,peak_amp,amp_error))

print("\n📊 Kết quả kiểm tra biên độ:")
print("Tần số đặt (MHz) | Biên độ (dBm) | Sai số (dBm)")
print("-" * 50)
for amp_set, amp_meas, amp_err in results_amp:
    print(f"{amp_set:<16} | {amp_meas:<16} | {amp_err:<16}")


# Kiểm tra chế độ AM máy phát
print("🔹 Kiểm tra AM.")
sig_anl.write("SYST:PRES:USER:ALL")
time.sleep(2)
sig_gen.write("SYST:PRES")
time.sleep(1)
am_data = [
		(2, 30),
		(2, 70),
		(249, 30),
		(249, 70),
		(920, 30),
		(920, 70),
		(2100, 30),
		(2100, 70),
		(3000, 30),
		(3000, 70)
]		
results_am = []
am_results_list = []  # Danh sách lưu kết quả đo AM

# Thực hiện đo cho từng bộ dữ liệu đã nhập
for freq_AM, depth_AM in am_data:
    sig_gen.write(f"FREQ {freq_AM} MHz")
    sig_gen.write(f"POW -10 dBm")
    sig_gen.write("AM:STAT ON")
    sig_gen.write("AM:SOUR INT")
    sig_gen.write("AM:INT:FREQ 1 kHz")
    sig_gen.write(f"AM:DEPTH {depth_AM}")
    sig_gen.write("OUTP ON")
    sig_gen.write("OUTP:MOD:STAT ON")
    time.sleep(1)

    sig_anl.write("INST:SEL ADEMOD")
    sig_anl.write("INIT:AM")
    time.sleep(1)

    sig_anl.write(f"FREQ:CENT {freq_AM} MHz")
    sig_anl.write("AM:FREQ:SPAN 15 kHz")
    sig_anl.write("AM:BAND 100 Hz")
    sig_anl.write("AM:BAND:CHAN 3 kHz")
    time.sleep(1)

    sig_anl.write("CALC:AM:MARK1:MAX")
    sig_anl.write("CALC:AM:MARK1:CENT")
    time.sleep(2)

    # Đọc kết quả đo từ máy phân tích tín hiệu
    am_result = sig_anl.query("READ:AM?")
    am_values = [float(val) for val in am_result.split(",")]

    # Lưu kết quả vào danh sách
    am_results_list.append((am_values))

# Hiển thị tất cả kết quả đã lưu
print("\n📊 Kết quả đo AM:")
for i, (am_values) in enumerate(am_results_list):
    print(f"  - Tần số trung tâm (MHz): {am_values[0] / 1_000_000}")
    print(f"  - Biên độ đo (dBm): {am_values[1]}")
    print(f"  - Độ sâu điều chế AM (%): {am_values[9]}")
    print(f"  - Tần số điều chế (Hz): {am_values[3]}")
    print(f"  - THD (%): {am_values[5]}")
    print(f"  - DISTORTION/TOTAL Vrms (%): {am_values[6]}")
    print("-" * 30)

# Kiểm tra chế độ FM máy phát
print("🔹 Kiểm tra FM.")
sig_anl.write("SYST:PRES:USER:ALL")
time.sleep(2)
sig_gen.write("SYST:PRES")
time.sleep(1)
fm_data = [
		(10, 5),
		(10, 10),
		(100, 5),
		(100, 10),
		(565, 5),
		(565, 10),
		(925, 5),
		(925, 10),
		(2750, 5),
		(2750, 10)
]
fm_results_list = [] 
for freq_FM, dev_FM in fm_data: 
 sig_gen.write("OUTP OFF")
 sig_gen.write("OUTP:MOD:STAT OFF")
 time.sleep(1)
 sig_gen.write(f"FREQ {freq_FM} MHz")
 sig_gen.write(f"POW -10 dBm") 
 sig_gen.write("OUTP ON")
 time.sleep(1)
 sig_anl.write("INST:SEL ADEMOD")
 sig_anl.write("INIT:FM")
 sig_anl.write(f"FREQ:CENT {freq_FM} MHz")
 time.sleep(1)
 sig_anl.write(f"FM:FREQ:SPAN {dev_FM * 5} kHz")
 sig_anl.write("FM:BAND 100 Hz")        
 sig_anl.write(f"FM:BAND:CHAN {dev_FM * 4} kHz")
 time.sleep(1)
 sig_anl.write("CALC:FM:MARK1:MAX")
 sig_anl.write("CALC:FM:MARK1:CENT")
 time.sleep(1)
 sig_gen.write("FM:STAT ON")
 sig_gen.write("FM:SOUR INT")
 sig_gen.write("FM:INT:FREQ 1 kHz")  
 sig_gen.write(f"FM:DEV {dev_FM} kHz")  
 sig_gen.write("OUTP:MOD:STAT ON")
 time.sleep(1)
 # Đọc kết quả đo từ máy phân tích tín hiệu
 fm_result = sig_anl.query("READ:FM?")
 time.sleep(1)
 fm_values = [float(val) for val in fm_result.split(",")]
 # Lưu kết quả vào danh sách fm_results_list
 fm_results_list.append((fm_values)) 

# Hiển thị kết quả
print("\n📊 Kết quả đo FM:")
for i, (fm_values) in enumerate(fm_results_list):
    print(f"  - Tần số trung tâm (MHz): {fm_values[0] / 1_000_000}")
    print(f"  - Biên độ đo (dBm): {fm_values[1]}")
    print(f"  - Độ dịch tần FM (kHz): {fm_values[9] / 1_000}")
    print(f"  - Tần số điều chế (Hz): {fm_values[3]}")
    print(f"  - THD (%): {fm_values[5]}")
    print(f"  - DISTORTION/TOTAL Vrms (%): {fm_values[6]}")
    print("-" * 30)

# In kết quả kiểm tra

# Lấy SN máy phát
SN_gen = sig_gen.query("*IDN?").split(",")[2]

# Lấy thời gian hiện tại
now = datetime.now()
timestamp = datetime.now().strftime("%m%d%Y_%H%M")
day = str(now.day)
month = str(now.month)
year = str(now.year)

# Đường dẫn file 
# Đường dẫn Desktop
desktop_path = os.path.expanduser("~/Desktop")

# Tạo thư mục KiemDinhMayDo trên Desktop
KiemDinhMayDo_folder = os.path.join(desktop_path, "KiemDinhMayDo")
os.makedirs(KiemDinhMayDo_folder, exist_ok=True)

# Tạo thư mục Template bên trong KiemDinhMayDo
template_folder = os.path.join(KiemDinhMayDo_folder, "Template")
os.makedirs(template_folder, exist_ok=True)

# Đường dẫn file template
template_path = os.path.join(template_folder, "kqkdN5181A.docx")

results_folder = os.path.join(KiemDinhMayDo_folder, "KetQuaKiemDinh")
os.makedirs(results_folder, exist_ok=True)
output_path = os.path.join(results_folder, f"kqkdN5181A_{SN_gen}_{timestamp}.docx")


# Mở file Word từ template
doc = Document(template_path)

# Kiểm tra xem file template có tồn tại không
if not os.path.exists(template_path):
    print("❌ Không tìm thấy file template")
    exit()

# Điền SN máy phát
placeholders = {
	"{SN_gen}": SN_gen,
	"{day}":day,
	"{month}":month,
	"{year}":year
}
for para in doc.paragraphs:
    for key, value in placeholders.items():
        if key in para.text:
            para.text = para.text.replace(key, value)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for key, value in placeholders.items():
                    if key in para.text:
                        para.text = para.text.replace(key, value)


# Điền vào bảng tần số
table = doc.tables[2]  # Chọn bảng
start_row = 2  # Hàng thứ 3 trong bảng (bỏ qua tiêu đề)

for i, freq in enumerate(results_freq):
    row_index = start_row + i  # Xác định hàng cần điền
    if row_index >= len(table.rows):  # Đảm bảo không vượt quá số hàng có sẵn
        break  

    row = table.rows[row_index].cells  # Lấy hàng cần điền
    row[2].text = f"{freq[1]/1000000}".replace('.', ',')  # Cột 3:

# Điền vào bảng biên độ
table = doc.tables[3]  
start_row = 2 

for i, amp in enumerate(results_amp):
    row_index = start_row + i  
    if row_index >= len(table.rows):  
        break  

    row = table.rows[row_index].cells  
    row[3].text = f"{amp[1]:.2f}".replace('.', ',')   

# Điền vào bảng méo AM
table = doc.tables[4] 
start_row = 1  

for i, dist in enumerate(am_results_list):
    row_index = start_row + i
    if row_index >= len(table.rows):  
        break  

    row = table.rows[row_index].cells  
    row[2].text = f"{dist[5]:.2f} %".replace('.', ',') 

# Điền vào bảng độ sâu điều chế AM
table = doc.tables[5]  
start_row = 1 

for i, depth in enumerate(am_results_list):
    row_index = start_row + i  
    if row_index >= len(table.rows):  
        break  

    row = table.rows[row_index].cells  
    row[3].text = f"{depth[9]:.2f} %".replace('.', ',')   

# Điền vào bảng độ sâu điều chế AM
table = doc.tables[6]  
start_row = 1 

for i, dev in enumerate(fm_results_list):
    row_index = start_row + i  
    if row_index >= len(table.rows):  
        break  

    row = table.rows[row_index].cells  
    row[3].text = f"{(dev[9] / 1000):.2f}".replace('.', ',')  

# Lưu file kết quả
doc.save(output_path)
print(f"✅ Kết quả đã lưu vào {output_path}")






